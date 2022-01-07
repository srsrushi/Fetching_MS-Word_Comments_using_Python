from docx import Document
from lxml import etree
import zipfile

ooXMLns = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

#Function to extract all the comments of document and returns a dictionary with comment id as key and comment string as value
def get_document_comments(docxFileName):
    comments_dict={}
    docxZip = zipfile.ZipFile(docxFileName)
    commentsXML = docxZip.read('word/comments.xml')
    et = etree.XML(commentsXML)
    comments = et.xpath('//w:comment',namespaces=ooXMLns)
    for c in comments:
        comment=c.xpath('string(.)',namespaces=ooXMLns)
        comment_id=c.xpath('@w:id',namespaces=ooXMLns)[0]
        comments_dict[comment_id]=comment
    return comments_dict

#Function to fetch all the comments in a paragraph
def paragraph_comments(paragraph,comments_dict):
    comments=[]
    for run in paragraph.runs:
        comment_reference=run._r.xpath("./w:commentReference")
        if comment_reference:
            comment_id=comment_reference[0].xpath('@w:id',namespaces=ooXMLns)[0]
            comment=comments_dict[comment_id]
            comments.append(comment)
    return comments


#Function to fetch all comments with their referenced table and return list like this [{'table text': [comment 1,comment 2]}]
def comments_with_reference_table(docxFileName):
    document = Document(docxFileName)
    tables = document.tables
    comments_dict=get_document_comments(docxFileName)
    comments_with_their_reference_table=[]
    contents_with_their_reference_table=[]
    table_names=[]
    table_comments_data = []
    table_count = 0

    for table in tables:
        i = 0
        table_count+=1
        for row in table.rows:
            j = 0
            row_text = []
            for cell in row.cells:
                for paragraph in table.cell(i, j).paragraphs:
                    row_text.append(paragraph.text)
                    if comments_dict:
                        comments=paragraph_comments(paragraph,comments_dict)  
                        if comments:
                            row_string = " , ".join(row_text)
                            table_name = "Table_{}".format(table_count)
                            table_comments_data.append({table_name:{row_string:comments}})
                j+=1
            i+=1  

    return table_comments_data

#Function to fetch all comments with their referenced paragraph and return list like this [{'Paragraph text': [comment 1,comment 2]}]
def comments_with_reference_paragraph(docxFileName):
    document = Document(docxFileName)
    comments_dict=get_document_comments(docxFileName)
    comments_with_their_reference_paragraph=[]
    for paragraph in document.paragraphs:  
        if comments_dict: 
            comments=paragraph_comments(paragraph,comments_dict)  
            if comments:
                comments_with_their_reference_paragraph.append({paragraph.text: comments})
    return comments_with_their_reference_paragraph



#filepath for the input document
document="test.docx" 

#get comments for paragraphs
print(comments_with_reference_paragraph(document))

#get comments for tables
print(comments_with_reference_table(document))
